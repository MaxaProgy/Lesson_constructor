# -*- coding: utf-8 -*-
def get_document_result_word(topic_lesson, teacher, subject, lesson_class, duration, competence, methods):
    from docx import Document

    document = Document()

    document.add_heading(f'Урок по теме "{topic_lesson}"', 0)

    p = document.add_paragraph('Учитель: ')
    p.add_run(f'{" ".join([item[0].upper() + item[1:].lower() for item in teacher.split()])}').bold = True
    p = document.add_paragraph('Предмет: ')
    p.add_run(f'{subject}').bold = True
    p = document.add_paragraph('Класс: ')
    p.add_run(f'{lesson_class}').bold = True
    p = document.add_paragraph('Длительность урока: ')
    p.add_run(f'{duration}').bold = True
    p = document.add_paragraph('Планируемое формирование компетенций: ')
    p.add_run(f'{", ".join(competence)}').bold = True

    document.add_heading('Этапы урока', level=1)

    table = document.add_table(rows=1, cols=5)
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = '№ этапа'
    hdr_cells[1].text = 'Название'
    hdr_cells[2].text = 'Тип'
    hdr_cells[3].text = 'Описание'
    hdr_cells[4].text = 'Длительность (минут)'
    for i in range(len(methods)):
        row_cells = table.add_row().cells
        row_cells[0].text = str(i + 1)
        row_cells[1].text = methods[i][1][0].upper() + methods[i][1][1:].lower()
        row_cells[2].text = methods[i][0]
        if methods[i][2] is None:
            row_cells[3].text = ' - '
        else:
            row_cells[3].text = methods[i][2]
        row_cells[4].text = methods[i][3]

    document.add_page_break()
    return document
